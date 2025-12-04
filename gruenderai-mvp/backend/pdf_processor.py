"""
PDF & DOCX Text Extraction
Handles both PDF and Word documents
"""

from io import BytesIO
import PyPDF2
from docx import Document

def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF bytes"""
    try:
        pdf_file = BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return text.strip()
    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")

def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX bytes"""
    try:
        docx_file = BytesIO(content)
        doc = Document(docx_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        
        return text.strip()
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")

def extract_text_from_file(content: bytes, content_type: str) -> str:
    """
    Extract text from file based on content type
    
    Args:
        content: File bytes
        content_type: MIME type
    
    Returns:
        Extracted text as string
    """
    
    if content_type == "application/pdf":
        return extract_text_from_pdf(content)
    
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(content)
    
    else:
        raise ValueError(f"Unsupported content type: {content_type}")

# Test function
if __name__ == "__main__":
    # Test with a sample file
    with open("sample.pdf", "rb") as f:
        content = f.read()
        text = extract_text_from_pdf(content)
        print(f"Extracted {len(text)} characters")
        print(text[:500])  # Print first 500 chars
